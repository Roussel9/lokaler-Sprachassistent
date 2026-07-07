from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/Praktikumsbericht_Lokale_KI_Anwendung.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 30, 45)
MUTED = RGBColor(90, 95, 105)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Seite ")
    run.font.size = Pt(9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_run(run, size=None, bold=None, italic=None, color=None, font="Calibri") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.text = "Praktikumsbericht - Lokale KI-Anwendung"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED
    add_page_number(footer)


def p(doc, text="", style=None, bold_prefix=None):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = para.add_run(bold_prefix)
        set_run(r1, bold=True)
        r2 = para.add_run(text[len(bold_prefix):])
        set_run(r2)
    else:
        run = para.add_run(text)
        set_run(run)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.167
    run = para.add_run(text)
    set_run(run)
    return para


def number(doc, text):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.167
    run = para.add_run(text)
    set_run(run)
    return para


def h(doc, level: int, text: str):
    return doc.add_heading(text, level=level)


def callout(doc, title: str, body: str, fill: str = CALLOUT) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.35])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(3)
    r = para.add_run(title)
    set_run(r, bold=True, color=DARK_BLUE)
    para2 = cell.add_paragraph()
    para2.paragraph_format.space_after = Pt(0)
    r2 = para2.add_run(body)
    set_run(r2)
    doc.add_paragraph()


def simple_table(doc, headers: list[str], rows: list[list[str]], widths: list[float], header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = ""
        set_cell_shading(hdr[i], header_fill)
        r = hdr[i].paragraphs[0].add_run(header)
        set_run(r, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(value)
            set_run(run, size=10.5)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_cover(doc: Document) -> None:
    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(80)
    r = p0.add_run("Praktikumsbericht / Expose")
    set_run(r, size=13, bold=True, color=MUTED)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    rt = title.add_run("Lokale KI-Anwendung: Entwicklung eines offlinefaehigen Sprachassistenten")
    set_run(rt, size=24, bold=True, color=DARK_BLUE)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(26)
    rs = sub.add_run("Von TTS-Qualitaetsanalyse und Optimierung bis zur synchronen Wake-Word-Sprachassistenz")
    set_run(rs, size=13, italic=True, color=MUTED)

    meta = [
        ["Praktikum", "Lokale KI-Anwendungen / Sprachassistenz"],
        ["Phase 1", "Recherche, Test und Optimierung von TTS, Voice Cloning und Sprachqualitaetsmetriken"],
        ["Phase 2", "Entwicklung eines eigenen lokalen KI-Sprachassistenten mit Wake Word und Echtzeit-Streaming"],
        ["Technologien", "Vosk, Whisper, Ollama/Gemma3, Piper, Gradio, SQLite, Python"],
        ["Stand", "Juli 2026"],
    ]
    simple_table(doc, ["Feld", "Inhalt"], meta, [1.45, 4.9], header_fill=LIGHT_BLUE)

    callout(
        doc,
        "Leitidee des Praktikums",
        "Das Praktikum untersucht, wie Sprache lokal verarbeitet, bewertet, optimiert und in eine nutzbare KI-Anwendung integriert werden kann. Der Schwerpunkt liegt auf Datenschutz, Offline-Faehigkeit, Echtzeitverhalten und praktischer Systemintegration.",
    )
    doc.add_page_break()


def add_toc_placeholder(doc: Document) -> None:
    h(doc, 1, "Inhaltsverzeichnis")
    p(doc, "Hinweis: In Microsoft Word kann dieses Inhaltsverzeichnis ueber Referenzen > Inhaltsverzeichnis automatisch aktualisiert werden.")
    for item in [
        "1 Einleitung",
        "2 Ausgangslage und Phase 1: TTS, Voice Cloning und Qualitaetsbewertung",
        "3 Phase 2: Entwicklung des lokalen Sprachassistenten",
        "4 Herkunft und Auswahl der Methoden",
        "5 Methoden, Metriken und Testaufbau",
        "6 Durchfuehrung, Ergebnisse und Diskussion",
        "7 Fazit",
        "8 Ausblick",
        "9 Anhang: Architektur, Testplan und Messvorlagen",
    ]:
        p(doc, item)
    doc.add_page_break()


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_toc_placeholder(doc)

    h(doc, 1, "1 Einleitung")
    h(doc, 2, "1.1 Problemstellung")
    p(doc, "Sprachassistenten sind im Alltag weit verbreitet, werden aber haeufig als Cloud-Dienste umgesetzt. Dadurch entstehen mehrere praktische und fachliche Probleme: Sprachdaten muessen ueber das Internet uebertragen werden, Kosten und Verfuegbarkeit haengen von externen APIs ab, und die Latenz der Antwort ist fuer den Nutzer nicht immer nachvollziehbar. Gleichzeitig wirkt die Ausgabe vieler Systeme unnatuerlich: Der Antworttext erscheint sofort, waehrend die Sprache erst nach Abschluss der Generierung erzeugt und abgespielt wird.")
    p(doc, "Im Praktikum wurde deshalb untersucht, wie eine lokale KI-Sprachanwendung aufgebaut werden kann, die Spracheingabe, Antwortgenerierung und Sprachausgabe ohne Cloud-Abhaengigkeit realisiert. Ein besonderes technisches Problem war die Synchronisation: Text und Audio sollten nicht getrennt erscheinen, sondern die Woerter sollten in der Benutzeroberflaeche genau dann sichtbar werden, wenn sie gesprochen werden.")
    h(doc, 2, "1.2 Motivation und Relevanz")
    p(doc, "Die Motivation ergibt sich aus drei Perspektiven. Erstens ist Datenschutz bei Sprachsystemen besonders wichtig, weil Sprache personenbezogene Informationen, Umgebungsgeraeusche und spontane Gedanken enthalten kann. Zweitens sind lokale KI-Systeme fuer Ausbildung und Forschung relevant, weil die gesamte Pipeline transparent analysiert und veraendert werden kann. Drittens zeigt das Projekt, dass moderne Sprach-KI nicht nur aus einem einzelnen Modell besteht, sondern aus einer Kette von Komponenten: Wake Word, Aufnahme, Speech-to-Text, LLM, Text-to-Speech, Audioausgabe, UI und Speicherung.")
    p(doc, "Die praktische Relevanz liegt darin, dass lokale Assistenten in Umgebungen eingesetzt werden koennen, in denen Internet, Datenschutz oder Kosten kritisch sind: Labore, Schulen, interne Unternehmenssysteme, medizinische oder pflegerische Assistenz, Smart-Home-Anwendungen und Barrierefreiheitswerkzeuge.")
    h(doc, 2, "1.3 Forschungsfrage")
    callout(
        doc,
        "Zentrale Forschungsfrage",
        "Wie laesst sich ein komplett lokaler, datenschutzfreundlicher KI-Sprachassistent entwickeln, der Wake-Word-Aktivierung, deutsche Spracherkennung, lokale Antwortgenerierung und synchrone Text-Sprach-Ausgabe in akzeptabler Qualitaet und Latenz bereitstellt?",
        fill=LIGHT_BLUE,
    )
    p(doc, "Daraus ergeben sich vier Teilfragen:")
    bullet(doc, "Welche lokalen Modelle und Werkzeuge eignen sich fuer Wake Word, Speech-to-Text, LLM und Text-to-Speech?")
    bullet(doc, "Wie kann die Qualitaet von synthetischer Sprache gemessen und optimiert werden?")
    bullet(doc, "Wie kann die Ausgabe so gestaltet werden, dass Text und Audio nicht auseinanderlaufen?")
    bullet(doc, "Welche Metriken und Tests zeigen, ob das System technisch funktioniert und fuer Nutzer verstaendlich ist?")
    h(doc, 2, "1.4 Abgrenzung")
    p(doc, "Der Bericht betrachtet kein kommerzielles Produkt, sondern einen funktionsfaehigen Prototyp im Praktikumskontext. Der Fokus liegt auf lokaler Ausfuehrung, Systemintegration und Messbarkeit. Nicht im Mittelpunkt stehen eine vollstaendige mobile App, Cloud-Skalierung, mehrsprachige Dialogfuehrung, Sicherheitshaertung fuer produktive Unternehmensumgebungen oder ein trainiertes eigenes Foundation Model. Auch Voice Cloning wurde in Phase 1 fachlich untersucht und getestet, aber die finale Anwendung nutzt aus Gruenden der Stabilitaet und Reproduzierbarkeit Piper mit deutscher Thorsten-Stimme.")

    h(doc, 1, "2 Ausgangslage und Phase 1: TTS, Voice Cloning und Qualitaetsbewertung")
    h(doc, 2, "2.1 Ziel der ersten Praktikumsphase")
    p(doc, "Vor der Entwicklung des eigenen Sprachassistenten lag der Schwerpunkt auf synthetischer Sprache: Welche TTS- und Voice-Cloning-Ansaetze gibt es, wie klingen sie, wie schnell laufen sie lokal und wie kann man die Qualitaet objektiv bewerten? Diese Phase war wichtig, weil ein Sprachassistent nicht nur eine richtige Antwort liefern muss. Die Antwort muss auch verstaendlich, angenehm und schnell genug gesprochen werden.")
    h(doc, 2, "2.2 SQUIM als Messmethode")
    p(doc, "SQUIM steht fuer Speech Quality and Intelligibility Measures. Es handelt sich um ein KI-Modell von Microsoft Research, das direkt in torchaudio verfuegbar ist. Das Besondere an SQUIM ist der reference-free bzw. non-intrusive Ansatz: Es wird keine saubere Originalaufnahme benoetigt. Das Modell bewertet nur die erzeugte Audiodatei und schaetzt daraus Qualitaets- und Verstaendlichkeitswerte.")
    simple_table(
        doc,
        ["Metrik", "Bedeutung", "Interpretation im Praktikum"],
        [
            ["PESQ", "Perceptual Evaluation of Speech Quality; urspruenglich aus der Telefonie.", "Skala ca. 1.0 bis 4.5. Werte unter 2.5 wirken messmethodisch schwach, koennen bei TTS aber auftreten, weil synthetische Sprache anders strukturiert ist als echte Mikrofonaufnahmen."],
            ["STOI", "Short-Time Objective Intelligibility; misst die Wortverstaendlichkeit.", "Skala 0.0 bis 1.0. Werte um 0.96 deuten auf sehr hohe Verstaendlichkeit hin."],
            ["SI-SDR", "Scale-Invariant Signal-to-Distortion Ratio; misst Verzerrungen im Signal.", "Hoeher ist besser. Nuetzlich, um Verzerrungen und Rauschen zu vergleichen."],
        ],
        [1.15, 2.35, 2.85],
        header_fill=LIGHT_BLUE,
    )
    p(doc, "Die SQUIM-Werte wurden nicht isoliert betrachtet, sondern zusammen mit Hoereindruck und Laufzeit. Besonders wichtig war die Erkenntnis, dass niedrige PESQ-Werte bei TTS nicht automatisch bedeuten, dass die Stimme unbrauchbar ist. SQUIM wurde vor allem auf echten Mikrofonaufnahmen mit Rauschen, Hall und Uebertragungsartefakten trainiert. Synthetische Stimmen koennen deshalb fuer das Modell ungewohnt wirken, obwohl sie fuer Menschen gut verstaendlich sind.")
    h(doc, 2, "2.3 Optimierung von TTS-Parametern")
    p(doc, "In der ersten Phase wurden Parameter untersucht, die die Qualitaet und Geschwindigkeit eines TTS-/Voice-Cloning-Systems beeinflussen. Die finale Konfiguration wurde nicht nur nach maximaler Qualitaet, sondern nach einem Kompromiss aus Verstaendlichkeit, Natuerlichkeit und Laufzeit gewaehlt.")
    simple_table(
        doc,
        ["Parameter", "Finale Wahl", "Begruendung"],
        [
            ["temperature", "0.8", "Stabiler Standardwert. Niedrigere Werte wirkten monotoner, hoehere Werte instabiler."],
            ["repetition_penalty", "5.0", "Bessere Qualitaet als 10.0; zu starke Wiederholungsstrafe kann die Generierung unnatuerlich machen."],
            ["top_p", "0.8", "Guter Kompromiss. top_p=0.6 erreichte teils bessere PESQ-Werte, war aber nicht die beste Wahl fuer Geschwindigkeit."],
            ["top_k", "30", "Standardwert blieb sinnvoll; top_k=15 und top_k=50 brachten keine klare Verbesserung."],
            ["diffusion_steps", "8", "Wichtigster Laufzeithebel. Die Reduktion von 25 auf 8 Schritte beschleunigte die Synthese massiv."],
            ["inference_cfg_rate", "0.0", "Kein festgestellter Qualitaetsverlust, aber weniger Rechenaufwand, weil CFG den Batch sonst verdoppelt."],
        ],
        [1.4, 1.2, 3.75],
        header_fill=LIGHT_BLUE,
    )
    p(doc, "Aus dieser Phase entstand eine methodische Grundlage fuer Phase 2: Sprachqualitaet muss immer mehrdimensional bewertet werden. Eine einzelne Zahl reicht nicht aus. Fuer eine lokale Anwendung zaehlen Qualitaet, Verstaendlichkeit, Reaktionszeit, Robustheit und Ressourcenverbrauch gemeinsam.")

    h(doc, 1, "3 Phase 2: Entwicklung des lokalen Sprachassistenten")
    h(doc, 2, "3.1 Zielsystem")
    p(doc, "In der zweiten Phase wurde eine eigene lokale KI-Anwendung entwickelt: ein deutscher Sprachassistent, der per Wake Word aktiviert wird, die Frage per Mikrofon aufnimmt, sie lokal transkribiert, eine Antwort mit einem lokalen LLM erzeugt, diese Antwort per Piper synthetisiert und gleichzeitig in einer Gradio-Oberflaeche wortweise anzeigt.")
    simple_table(
        doc,
        ["Komponente", "Auswahl", "Funktion im System"],
        [
            ["Wake Word", "Vosk Grammar Mode", "Aktivierung durch 'Computer' und nahe Varianten, ohne dauerhaft vollstaendige Sprache zu transkribieren."],
            ["STT", "Whisper / Faster-Whisper", "Deutsche Spracherkennung fuer die eigentliche Nutzerfrage."],
            ["LLM", "Gemma3 ueber Ollama", "Lokale Antwortgenerierung mit Streaming-Ausgabe."],
            ["TTS", "Piper Thorsten", "Lokale deutsche Sprachausgabe, reproduzierbar und schnell genug fuer Streaming."],
            ["UI", "Gradio", "Status, erkannte Frage, synchrone Antwortanzeige, Audioausgabe und Verlauf."],
            ["Speicherung", "SQLite", "Lokale Speicherung von Frage, Antwort und Laufzeiten fuer Auswertung."],
        ],
        [1.25, 1.75, 3.35],
        header_fill=LIGHT_BLUE,
    )
    h(doc, 2, "3.2 Architektur der Pipeline")
    p(doc, "Die Anwendung ist als Pipeline aufgebaut. Jede Komponente uebergibt ihr Ergebnis an die naechste Stufe. Der entscheidende Unterschied zu einer einfachen Batch-Loesung ist, dass die LLM-Antwort nicht erst komplett abgewartet wird. Stattdessen werden Token inkrementell verarbeitet, zu sprechbaren Saetzen gebuendelt, parallel synthetisiert und abgespielt.")
    code = (
        "Wake Word (Vosk) -> Mikrofonaufnahme -> Whisper STT -> Gemma3/Ollama Streaming\n"
        "                         -> SentenceChunker -> text_queue -> PiperWorker\n"
        "                         -> audio_queue -> AudioPlayer -> UI + Audioausgabe\n"
        "                         -> SQLite: Frage, Antwort, STT/LLM/TTS/Gesamtzeit"
    )
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.2)
    para.paragraph_format.space_after = Pt(10)
    r = para.add_run(code)
    set_run(r, font="Consolas", size=9.5, color=INK)
    h(doc, 2, "3.3 Synchronisierte Text- und Audioausgabe")
    p(doc, "Das Kernproblem der zweiten Phase war die Synchronisation. In vielen Sprachsystemen wird zuerst der komplette Text generiert, danach synthetisiert und anschliessend abgespielt. Dadurch sieht der Nutzer den Text frueher als er ihn hoert. In diesem Projekt wurde die Ausgabe so gebaut, dass der AudioPlayer die bereits abgespielten Samples als Fortschritt nutzt und daraus berechnet, wie viele Woerter des aktuellen Satzes angezeigt werden duerfen.")
    p(doc, "Die Pipeline arbeitet mit zwei begrenzten Queues: einer Text-Queue und einer Audio-Queue. Der SentenceChunker sammelt den LLM-Tokenstrom zu Saetzen oder sinnvollen Teilstuecken. Der PiperWorker synthetisiert diese Stuecke im Hintergrund. Der AudioPlayer spielt sie ab und meldet der UI erst dann neue Woerter, wenn der entsprechende Anteil des Audios wirklich geschrieben wurde. Dadurch entsteht eine natuerliche Wort-fuer-Wort-Anzeige synchron zur Stimme.")
    callout(
        doc,
        "Technischer Kern",
        "Die Anwendung nutzt Backpressure: Wenn TTS oder Audio langsamer sind als das LLM, blockieren die Queues. Dadurch laeuft die Pipeline nicht unkontrolliert voraus, sondern passt sich automatisch an die langsamste Komponente an.",
        fill=LIGHT_BLUE,
    )
    h(doc, 2, "3.4 Wake-Word-Entscheidung")
    p(doc, "Fuer das Wake Word wurde Vosk statt Whisper verwendet. Whisper ist fuer laengere Transkription sehr stark, aber als Wake-Word-System zu schwerfaellig und zu offen: Bei kurzen oder verrauschten Signalen kann ein generatives STT-Modell falsche Woerter halluzinieren. Vosk kann dagegen mit einer kleinen Grammatik betrieben werden. Die Erkennung wird damit auf wenige erlaubte Phrasen begrenzt, zum Beispiel 'computer', 'komputer' oder 'hey computer'.")
    p(doc, "Im Projekt wurde die Wake-Word-Logik zusaetzlich verschaerft: Es werden nur finale Vosk-Ergebnisse akzeptiert, keine optimistischen Partial-Results. Zudem werden zu kurze Impulse, lange Sprache, zu leises Rauschen und uebersteuertes Audio verworfen. Das ist wichtig, weil ein Wake Word nicht 'alles verstehen' soll, sondern nur einen sehr kleinen Ausloesewortschatz robust erkennen muss.")

    h(doc, 1, "4 Herkunft und Auswahl der Methoden")
    h(doc, 2, "4.1 Moegliche Alternativen")
    simple_table(
        doc,
        ["Aufgabe", "Moegliche Methoden", "Gewaehlte Methode und Grund"],
        [
            ["Wake Word", "Porcupine, Snowboy, Whisper, Vosk, eigenes kleines Modell", "Vosk, weil offline, lokal, leichtgewichtig und mit Grammar Mode auf wenige Woerter begrenzbar."],
            ["STT", "Whisper, Vosk Volltranskription, cloudbasierte APIs", "Whisper/Faster-Whisper, weil die deutsche Transkription im Vergleich zu kleinen lokalen STT-Loesungen genauer ist."],
            ["LLM", "Cloud-LLM, lokale llama.cpp-Modelle, Ollama", "Gemma3 ueber Ollama, weil lokale Ausfuehrung und einfache Streaming-Schnittstelle moeglich sind."],
            ["TTS", "Piper, Coqui/XTTS, Bark, kommerzielle TTS-APIs", "Piper, weil lokal, stabil, relativ schnell und mit deutscher Thorsten-Stimme gut reproduzierbar."],
            ["UI", "CLI, Desktop-App, Web-Frontend, Gradio", "Gradio, weil schnell prototypisierbar und geeignet fuer Status, Text, Audio und Verlauf."],
            ["Bewertung", "Subjektive Hoertests, SQUIM, RTF, Latenz, WER", "Kombination, weil Sprachsysteme nicht mit einer einzelnen Metrik ausreichend beschrieben werden koennen."],
        ],
        [1.15, 2.25, 2.95],
        header_fill=LIGHT_BLUE,
    )
    h(doc, 2, "4.2 Begruendung der Auswahl")
    p(doc, "Die Auswahl folgt dem Ziel des Praktikums: eine lokale, nachvollziehbare und praktisch lauffaehige Anwendung zu entwickeln. Cloud-Dienste waeren oft qualitativ stark, widersprechen aber dem Datenschutz- und Offline-Ziel. Ein selbst trainiertes Wake-Word-Modell waere wissenschaftlich interessant, haette aber den Rahmen des Praktikums gesprengt. Deshalb wurden etablierte lokale Komponenten kombiniert und dort angepasst, wo die Systemintegration es erforderte.")
    p(doc, "Wichtig ist dabei die Trennung der Aufgaben: Vosk wird nicht fuer komplexe Fragen verwendet, sondern nur fuer das kleine Aktivierungsproblem. Whisper wird nicht fuer dauerhaftes Lauschen verwendet, sondern fuer die eigentliche Frage nach der Aktivierung. Piper ist nicht das experimentellste TTS-System, aber fuer die finale Anwendung stabiler und leichter kontrollierbar als ein groesseres Voice-Cloning-Modell.")

    h(doc, 1, "5 Methoden, Metriken und Testaufbau")
    h(doc, 2, "5.1 Methodischer Ansatz")
    p(doc, "Das Praktikum kombiniert explorative Forschung, Prototyping und systematische Evaluation. In Phase 1 wurden TTS-Systeme und Parameter getestet. In Phase 2 wurde aus den Erkenntnissen eine modulare Anwendung gebaut. Die Evaluation betrachtet nicht nur Modellqualitaet, sondern auch das Zusammenspiel der Komponenten.")
    number(doc, "Recherche und Auswahl lokaler Modelle fuer Sprache und LLM.")
    number(doc, "Implementierung einer modularen Pipeline mit klaren Schnittstellen.")
    number(doc, "Messung von Qualitaet, Latenz und Robustheit.")
    number(doc, "Iterative Optimierung, z.B. Wake-Word-Striktheit und TTS-Streaming.")
    number(doc, "Diskussion der Grenzen und Ableitung weiterer Verbesserungen.")
    h(doc, 2, "5.2 Metriken")
    simple_table(
        doc,
        ["Bereich", "Metrik", "Was wird gemessen?", "Bewertung"],
        [
            ["TTS", "STOI", "Verstaendlichkeit der synthetischen Stimme.", "Hoeher ist besser; Werte um 0.96 sind sehr gut."],
            ["TTS", "PESQ", "Wahrgenommene Sprachqualitaet.", "Mit Vorsicht interpretieren, da TTS nicht wie echte Mikrofonaufnahmen verteilt ist."],
            ["TTS", "SI-SDR", "Verzerrung im Audiosignal.", "Hoeher ist besser."],
            ["TTS/LLM", "RTF", "Real-Time-Factor: Synthesezeit im Verhaeltnis zur Audiolaenge.", "RTF < 1 waere Echtzeit; hoehere Werte zeigen Optimierungsbedarf."],
            ["System", "STT-Zeit", "Dauer der Transkription.", "Wichtig fuer wahrgenommene Reaktionszeit."],
            ["System", "LLM-Zeit / Tokens pro Sekunde", "Antwortgenerierung und Streaming-Leistung.", "Zeigt, ob lokale Hardware ausreichend ist."],
            ["System", "TTS-Zeit", "Zeit fuer Synthese der Antwortsaetze.", "Entscheidend fuer synchrone Ausgabe."],
            ["Wake Word", "False Positives / False Negatives", "Falsches Ausloesen bzw. verpasste Aktivierung.", "Wake Word muss lieber strenger als zu offen sein."],
            ["STT", "WER/CER", "Wort- bzw. Zeichenfehlerrate der Transkription.", "Optional mit manuell referenzierten Testsaetzen."],
            ["UI", "Sync-Fehler", "Abweichung zwischen gehoertem Wort und sichtbarem Wort.", "Subjektiv und optional per Logging/Video messbar."],
        ],
        [0.85, 1.1, 2.65, 1.75],
        header_fill=LIGHT_BLUE,
    )
    h(doc, 2, "5.3 Testaufbau")
    p(doc, "Der Testaufbau besteht aus einem lokalen Rechner mit Mikrofon, Python-Umgebung, lokalen Modellen und Gradio-UI. Ollama stellt das LLM lokal bereit. Die Anwendung speichert fuer jede Interaktion die Frage, Antwort sowie STT-, LLM-, TTS- und Gesamtzeit in SQLite. Dadurch lassen sich Messwerte spaeter aus dem Verlauf auswerten.")
    simple_table(
        doc,
        ["Testfall", "Durchfuehrung", "Erwartetes Ergebnis"],
        [
            ["Wake Word positiv", "Mehrfach 'Computer', 'hey computer' und nahe deutsche Varianten sprechen.", "Assistent startet Aufnahme zuverlaessig."],
            ["Wake Word negativ", "Normale Saetze, Fremdsprache, Hintergrundsprache und Computergeraeusche abspielen.", "Keine Aktivierung, wenn kein Wake Word enthalten ist."],
            ["STT", "Definierte deutsche Testsaetze einsprechen und mit Referenztext vergleichen.", "Niedrige WER/CER, besonders bei klarer Sprache."],
            ["LLM", "Standardfragen, Rechenfragen und Kontextfragen stellen.", "Kurze, deutsche, passende Antworten; Streaming aktiv."],
            ["TTS", "Antworten verschiedener Laenge erzeugen.", "Verstaendlich, nicht uebersteuert, keine starken Pausenfehler."],
            ["Synchronisation", "Antwort im UI beobachten und Audio parallel hoeren.", "Woerter erscheinen ungefaehr dann, wenn sie gesprochen werden."],
            ["End-to-End", "Mehrere komplette Dialoge fuehren.", "System kehrt nach jeder Antwort in den Idle-Modus zurueck."],
        ],
        [1.25, 3.0, 2.1],
        header_fill=LIGHT_BLUE,
    )

    h(doc, 1, "6 Durchfuehrung, Ergebnisse und Diskussion")
    h(doc, 2, "6.1 Durchfuehrung der ersten Phase")
    p(doc, "In Phase 1 wurden TTS- und Voice-Cloning-Ansaetze betrachtet und mit SQUIM sowie Laufzeitmessungen bewertet. Die wichtigsten Ergebnisse waren: Die Verstaendlichkeit synthetischer Sprache kann sehr hoch sein, auch wenn PESQ nicht immer hoch ausfaellt. Gleichzeitig zeigte sich, dass Diffusionsschritte und CFG-Parameter enorme Auswirkungen auf die Laufzeit haben. Die finale Parameterwahl priorisierte daher einen praxistauglichen Kompromiss.")
    p(doc, "Die Erkenntnis aus Phase 1 war fuer die finale Anwendung entscheidend: Ein Sprachassistent braucht nicht die theoretisch beste Stimme, sondern eine stabile, lokale, verstaendliche und schnell genug synthetisierbare Stimme. Deshalb wurde fuer Phase 2 Piper gewaehlt.")
    h(doc, 2, "6.2 Durchfuehrung der zweiten Phase")
    p(doc, "Die zweite Phase begann mit dem Aufbau der Grundpipeline: Wake Word, Aufnahme, Whisper-Transkription, Ollama-Anfrage, Piper-Ausgabe und Gradio-UI. Danach wurde die Anwendung schrittweise verbessert. Besonders wichtig waren die Thread-Sicherheit des UI-Zustands, das Freigeben des Mikrofons waehrend Aufnahme und TTS, die lokale Speicherung der Laufzeiten und die Umstellung auf eine echte Streaming-Pipeline.")
    p(doc, "Die finale Pipeline erzeugt pro Antwort eine eigene StreamingPipeline. Der LLM-Stream wird ueber einen Callback in den SentenceChunker geleitet. Vollstaendige Saetze oder sinnvolle Teilstuecke werden in eine begrenzte Text-Queue gelegt. Der PiperWorker synthetisiert daraus AudioTasks. Der AudioPlayer spielt die Audiodaten in Bloecken ab, schreibt optional WAV-Dateien und zeigt in der UI nur die Woerter an, deren Audiofortschritt bereits erreicht wurde.")
    h(doc, 2, "6.3 Ergebnisse")
    bullet(doc, "Ein komplett lokaler Prototyp wurde umgesetzt: Wake Word, STT, LLM, TTS, UI und Datenbank laufen lokal.")
    bullet(doc, "Die Antwortausgabe ist nicht nur textuell gestreamt, sondern mit der Audioausgabe gekoppelt.")
    bullet(doc, "Die SQLite-Datenbank speichert neben Frage und Antwort auch Laufzeiten, was spaetere Messreihen ermoeglicht.")
    bullet(doc, "Die Wake-Word-Erkennung wurde von einer zu offenen Erkennung auf eine eng begrenzte Computer-Grammatik umgestellt.")
    bullet(doc, "Die Systemarchitektur ist modular genug, um einzelne Komponenten spaeter auszutauschen, z.B. ein anderes LLM oder eine andere TTS-Stimme.")
    h(doc, 2, "6.4 Diskussion")
    p(doc, "Das Projekt zeigt, dass ein lokaler Sprachassistent mit heutigen Open-Source-Komponenten realistisch umsetzbar ist. Die groesste technische Herausforderung liegt weniger in einem einzelnen Modell, sondern im Zusammenspiel der Komponenten. Sobald STT, LLM und TTS getrennt arbeiten, entstehen Synchronisations-, Latenz- und Ressourcenprobleme.")
    p(doc, "Die Entscheidung fuer lokale Verarbeitung bringt klare Vorteile: Datenschutz, keine API-Kosten und technische Kontrolle. Gleichzeitig entstehen Nachteile: Die Qualitaet und Geschwindigkeit haengen stark von der lokalen Hardware ab. Groessere Modelle verbessern Antworten oder Transkription, koennen aber Latenz und Speicherbedarf erhoehen. Die Anwendung muss daher immer als Kompromiss aus Qualitaet, Geschwindigkeit und Robustheit bewertet werden.")
    p(doc, "Bei der Wake-Word-Erkennung wurde deutlich, dass ein Wake Word nicht wie normale Spracherkennung behandelt werden darf. Ein zu tolerantes System ist im Alltag stoerend, weil es bei beliebiger Sprache ausloest. Deshalb ist eine engere Grammatik, das Ignorieren von Partial-Results und die Audio-Plausibilitaetspruefung fachlich sinnvoll.")
    h(doc, 2, "6.5 Gueltigkeitsbereich der Ergebnisse")
    p(doc, "Die Ergebnisse gelten fuer den Praktikumskontext und fuer eine lokale Einzelplatzanwendung. Sie zeigen, dass die Architektur funktioniert und dass lokale Komponenten sinnvoll kombiniert werden koennen. Sie beweisen noch nicht, dass das System unter allen Umgebungsbedingungen, Mikrofonen, Dialekten und Hardwarekonfigurationen gleich robust ist. Fuer eine produktive Bewertung waeren groessere Testreihen mit verschiedenen Sprechern, Geraeuschsituationen und Hardwareprofilen notwendig.")

    h(doc, 1, "7 Fazit")
    p(doc, "Das Praktikum entwickelte sich von einer Untersuchung synthetischer Sprache zu einer vollstaendigen lokalen KI-Sprachanwendung. Phase 1 lieferte das Verstaendnis fuer TTS-Qualitaet, Metriken und Laufzeitoptimierung. Phase 2 setzte dieses Wissen in einem eigenen Sprachassistenten um.")
    p(doc, "Die zentrale Forschungsfrage kann grundsaetzlich positiv beantwortet werden: Ein lokaler, datenschutzfreundlicher KI-Sprachassistent mit Wake Word, deutscher Spracherkennung, lokaler Antwortgenerierung und synchroner Text-Sprach-Ausgabe ist umsetzbar. Der Prototyp zeigt die technische Machbarkeit und macht gleichzeitig sichtbar, wo die Grenzen liegen: Wake-Word-Robustheit, Latenz, Hardwarelast und die Bewertung synthetischer Sprache muessen sorgfaeltig behandelt werden.")
    p(doc, "Besonders wichtig ist die methodische Einsicht, dass Tests und Metriken nicht nur am Ende stehen. Sie beeinflussen bereits die Architekturentscheidung. SQUIM, STOI, PESQ, RTF, Laufzeitmessungen und Wake-Word-Fehlerraten helfen dabei, technische Entscheidungen zu begruenden und nicht nur subjektiv zu treffen.")

    h(doc, 1, "8 Ausblick")
    bullet(doc, "Groessere Testreihe mit mehreren Sprechern, Dialekten, Mikrofonabstaenden und Hintergrundgeraeuschen.")
    bullet(doc, "Automatische Auswertung der SQLite-Verlaufsdaten als CSV oder Dashboard mit STT-, LLM-, TTS- und Gesamtzeiten.")
    bullet(doc, "Praezisere Messung der Synchronisation, z.B. ueber Zeitstempel pro Wort oder Videoanalyse der UI.")
    bullet(doc, "Vergleich verschiedener lokaler LLMs hinsichtlich Antwortqualitaet, Tokens pro Sekunde und RAM-Verbrauch.")
    bullet(doc, "Optionaler Wechsel zwischen Piper-Stimmen oder Integration eines optimierten Voice-Cloning-Modells, falls Qualitaet und Geschwindigkeit ausreichen.")
    bullet(doc, "Verbesserung der Wake-Word-Erkennung durch Sprecher-/Umgebungsprofile, Kalibrierung der Audio-Schwellen und systematische False-Positive-Tests.")
    bullet(doc, "Erweiterung der UI um Testmodus, Messprotokoll und Exportfunktion fuer Praktikumsdaten.")

    h(doc, 1, "9 Anhang: Messvorlagen")
    h(doc, 2, "9.1 Vorlage fuer End-to-End-Messungen")
    simple_table(
        doc,
        ["Nr.", "Frage", "STT s", "LLM s", "TTS s", "Gesamt s", "Kommentar"],
        [
            ["1", "Wie spaet ist es?", "", "", "", "", ""],
            ["2", "Erklaere kurz, was ein lokaler Sprachassistent ist.", "", "", "", "", ""],
            ["3", "Was ist zwei plus fuenf?", "", "", "", "", ""],
            ["4", "Fasse den letzten Dialog kurz zusammen.", "", "", "", "", ""],
        ],
        [0.45, 2.35, 0.7, 0.7, 0.7, 0.8, 0.65],
        header_fill=LIGHT_BLUE,
    )
    h(doc, 2, "9.2 Vorlage fuer Wake-Word-Test")
    simple_table(
        doc,
        ["Testblock", "Material", "Erwartung", "Treffer", "Fehler"],
        [
            ["Positiv", "20x 'Computer' in normaler Lautstaerke", "Hohe Trefferquote", "", ""],
            ["Leise", "10x 'Computer' leise gesprochen", "Moeglichst robuste Treffer", "", ""],
            ["Negativ", "5 Minuten normale deutsche Sprache ohne Wake Word", "0 False Positives", "", ""],
            ["Fremdsprache", "Englische/franzoesische Saetze ohne Wake Word", "0 False Positives", "", ""],
            ["Stoerung", "Tastatur, Luefter, TTS-Audio, Raumgeraeusche", "0 False Positives", "", ""],
        ],
        [1.0, 2.25, 1.45, 0.8, 0.85],
        header_fill=LIGHT_BLUE,
    )
    h(doc, 2, "9.3 Kurzbewertung fuer die Praesentation")
    callout(
        doc,
        "Kernaussage fuer das Expose",
        "Das Praktikum zeigt den Weg von der Bewertung einzelner Sprachmodelle zur Integration einer kompletten lokalen KI-Anwendung. Der wichtigste technische Beitrag ist die synchrone Echtzeit-Ausgabe: Die KI generiert, Piper spricht, und die UI zeigt die Woerter genau entlang des Audiostreams.",
        fill=LIGHT_BLUE,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
